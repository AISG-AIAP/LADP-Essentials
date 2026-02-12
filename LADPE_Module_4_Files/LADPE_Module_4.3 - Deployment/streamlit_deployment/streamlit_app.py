import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime

# API Configuration

#######################amend this part of the code to use the Flowise API endpoint########################
import requests

API_URL = "http://localhost:3000/api/v1/prediction/e2a3465a-21b2-41c0-a64f-957cd3bdfda2"

def query(payload):
    response = requests.post(API_URL, json=payload)
    return response.json()

##########################################################################################################

def create_word_document(topic, report_text):
    """Create a Word document from the report text"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Business Report: {topic}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    timestamp_run = timestamp.runs[0]
    timestamp_run.font.size = Pt(10)
    timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Add spacing
    
    # Add the report content
    # Split by lines and preserve formatting
    for line in report_text.split('\n'):
        if line.strip().startswith('# '):
            doc.add_heading(line.replace('# ', ''), 1)
        elif line.strip().startswith('## '):
            doc.add_heading(line.replace('## ', ''), 2)
        elif line.strip().startswith('### '):
            doc.add_heading(line.replace('### ', ''), 3)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip():
            doc.add_paragraph(line.strip())
    
    return doc

def extract_report_text(output):
    """Extract the report text from the API output"""
    if not output:
        return None
    
    # Try to get the main text from the response
    if 'text' in output:
        return output['text']
    
    # Try to get it from agentFlowExecutedData
    if 'agentFlowExecutedData' in output:
        for node in output['agentFlowExecutedData']:
            if node.get('nodeLabel') == 'Report Writing agent':
                if 'data' in node and 'output' in node['data'] and 'content' in node['data']['output']:
                    return node['data']['output']['content']
    
    return None

# Streamlit App
def main():
    st.set_page_config(
        page_title="Business Report Generator",
        layout="wide"
    )
    
    st.title("Business Report Generator")
    st.markdown("---")
    
    # Sidebar for information
    with st.sidebar:
        st.header("About")
        st.info(
            "This app generates comprehensive business reports on any topic using AI agents. "
            "Simply enter a topic, and our multi-agent system will research, analyze, and write "
            "a professional report for you."
        )
        st.markdown("### Features:")
        st.markdown("- Research Agent")
        st.markdown("- Analyst Agent")
        st.markdown("- Report Writing Agent")
        st.markdown("- Save to Word Document")
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("Enter Your Topic")
        topic = st.text_input(
            "Topic",
            placeholder="e.g., Autonomous Vehicles, Renewable Energy, Artificial Intelligence...",
            label_visibility="collapsed"
        )
    
    with col2:
        st.write("")  # Spacing
        st.write("")  # Spacing
        generate_button = st.button("Generate Report", type="primary", use_container_width=True)
    
    # Initialize session state
    if 'report_data' not in st.session_state:
        st.session_state.report_data = None
    if 'current_topic' not in st.session_state:
        st.session_state.current_topic = None
    
    # Generate report when button is clicked
    if generate_button and topic:
        with st.spinner(f"Generating report on '{topic}'... This may take a moment."):
            output = query({"question": topic})
            if output:
                report_text = extract_report_text(output)
                if report_text:
                    st.session_state.report_data = {
                        'topic': topic,
                        'report_text': report_text,
                        'full_output': output
                    }
                    st.session_state.current_topic = topic
                    st.success("Report generated successfully!")
                else:
                    st.error("Could not extract report text from the response.")
    elif generate_button and not topic:
        st.warning("Please enter a topic first.")
    
    # Display report if available
    if st.session_state.report_data:
        st.markdown("---")
        st.subheader(f"Report: {st.session_state.report_data['topic']}")
        
        # Create tabs for different views
        tab1, tab2 = st.tabs(["Formatted Report", "Raw Data"])
        
        with tab1:
            # Display the formatted report
            st.markdown(st.session_state.report_data['report_text'])
            
            # Download button
            st.markdown("---")
            col1, col2, col3 = st.columns([1, 1, 2])
            
            with col1:
                # Save to Word document
                doc = create_word_document(
                    st.session_state.report_data['topic'],
                    st.session_state.report_data['report_text']
                )
                
                # Save to buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                filename = f"report_{st.session_state.report_data['topic'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                
                st.download_button(
                    label="Download Word Document",
                    data=buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col2:
                # Download as text
                text_filename = f"report_{st.session_state.report_data['topic'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                st.download_button(
                    label="Download Text File",
                    data=st.session_state.report_data['report_text'],
                    file_name=text_filename,
                    mime="text/plain",
                    use_container_width=True
                )
        
        with tab2:
            # Display raw JSON data
            st.json(st.session_state.report_data['full_output'])
    
    # Footer
    st.markdown("---")
    st.markdown(
        "<div style='text-align: center; color: gray;'>"
        "Powered by Flowise AI | Multi-Agent Report Generation System"
        "</div>",
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()

